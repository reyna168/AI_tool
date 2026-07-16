"""多格式文件載入器"""

from dataclasses import dataclass
from pathlib import Path


@dataclass
class Document:
    """文件物件"""

    content: str
    metadata: dict

    def __repr__(self):
        source = self.metadata.get("source", "unknown")
        return f"Document(source='{source}', length={len(self.content)})"


def load_txt(path: Path) -> Document:
    content = path.read_text(encoding="utf-8")
    return Document(content=content, metadata={"source": str(path), "type": "txt"})


def load_md(path: Path) -> Document:
    content = path.read_text(encoding="utf-8")
    return Document(content=content, metadata={"source": str(path), "type": "markdown"})


def load_pdf(path: Path) -> Document:
    try:
        import pypdf

        reader = pypdf.PdfReader(str(path))
        content = "\n".join(page.extract_text() or "" for page in reader.pages)
        return Document(content=content, metadata={"source": str(path), "type": "pdf"})
    except ImportError:
        raise ImportError("需要安裝 pypdf: pip install pypdf")


def load_csv(path: Path) -> Document:
    import csv

    with open(path, "r", encoding="utf-8") as f:
        reader = csv.reader(f)
        lines = [", ".join(row) for row in reader]
    content = "\n".join(lines)
    return Document(content=content, metadata={"source": str(path), "type": "csv"})


def load_json(path: Path) -> Document:
    import json

    data = path.read_text(encoding="utf-8")
    parsed = json.loads(data)
    content = json.dumps(parsed, ensure_ascii=False, indent=2)
    return Document(content=content, metadata={"source": str(path), "type": "json"})


def load_python(path: Path) -> Document:
    content = path.read_text(encoding="utf-8")
    return Document(content=content, metadata={"source": str(path), "type": "python"})


def load_docx(path: Path) -> Document:
    try:
        import docx

        doc = docx.Document(str(path))
        content = "\n".join(para.text for para in doc.paragraphs)
        return Document(content=content, metadata={"source": str(path), "type": "docx"})
    except ImportError:
        raise ImportError("需要安裝 python-docx: pip install python-docx")


LOADERS: dict[str, callable] = {
    ".txt": load_txt,
    ".md": load_md,
    ".pdf": load_pdf,
    ".csv": load_csv,
    ".json": load_json,
    ".py": load_python,
    ".docx": load_docx,
}


def load_document(path: str | Path) -> Document:
    """載入單一文件"""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"找不到檔案: {path}")

    suffix = path.suffix.lower()
    loader = LOADERS.get(suffix)
    if not loader:
        raise ValueError(f"不支援的檔案格式: {suffix}")

    return loader(path)


def load_directory(
    dir_path: str | Path, extensions: list[str] | None = None
) -> list[Document]:
    """遞迴載入目錄下所有支援的文件"""
    dir_path = Path(dir_path)
    if not dir_path.is_dir():
        raise NotADirectoryError(f"不是有效的目錄: {dir_path}")

    if extensions is None:
        extensions = list(LOADERS.keys())

    docs = []
    for ext in extensions:
        for file_path in dir_path.rglob(f"*{ext}"):
            try:
                doc = load_document(file_path)
                if doc.content.strip():
                    docs.append(doc)
            except Exception as e:
                print(f"[WARN] 載入 {file_path} 失敗: {e}")

    return docs
